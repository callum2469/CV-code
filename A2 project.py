from tkinter import *
from tkinter import ttk
import csv
import sys
import random
from datetime import date, timedelta
from tkinter import messagebox
import tkinter.simpledialog
import datetime
import time
import math
from datetime import datetime
import tkinter as tk
import mysql.connector
from tkcalendar import Calendar
from tkcalendar import Calendar,DateEntry
from PIL import Image, ImageTk
from datetime import date
import keyboard
import matplotlib.pyplot as plt
import smtplib, ssl
import boto3
from docx import Document
from docx.shared import Inches
from varname import nameof
import xlsxwriter
import excel2img
from docx.shared import Inches, Cm
import easygui
import os
import folium
import webbrowser
####################################### mySQL ################################################
mydb = mysql.connector.connect(
        host="localhost",
        user="root",
        password="btec12345",
        database="bayburn_fitness_database"
        )
######################################## goto functions #######################################
def manager_add_goto():
    raise_frame(manager_add)

def manager_info_goto():
    raise_frame(manager_info)

def manager_report_goto():
    update_treeview14()
    raise_frame(manager_report)

def trainer_make_classes_goto():
    raise_frame(trainer_make_classes)

def trainer_info_goto():
    raise_frame(trainer_info)

def trainer_see_bookings_goto():
    raise_frame(trainer_see_bookings)

def client_book_goto():
    raise_frame(client_book)
    refresh_bookings()
    refresh_bookings()
    refresh_alerts()
    alert_update2()

def client_info_goto():
    raise_frame(client_info)
    alert_update()
    refresh_alerts2()
    update_treeview13()
    

def client_track_goto():
    update2()
    raise_frame(client_track)
    alert_update()
    alert_update2()

def logout():
    password1.set("")
    username1.set("")
    raise_frame(login)
    alert_update()
    alert_update2()
    

######################################## hover button #########################################

class HoverButton(tk.Button):
    def __init__(self, master, **kw):
        tk.Button.__init__(self,master=master,**kw)
        self.defaultBackground = self["background"]
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)

    def on_enter(self, e):
        self['background'] = self['activebackground']

    def on_leave(self, e):
        self['background'] = self.defaultBackground
######################################## functions ############################################
def raise_frame(frame):
    frame.tkraise()

#clears all entry fields in the add user screen
def clear_all():
    height.set("")
    weight.set("")
    username_new.set("")
    password_new.set("")
    firstname.set("")
    lastname.set("")
    phone.set("")
    email.set("")
    postcode.set("")
    address.set("")
    text1.delete('1.0', END)
    text2.delete('1.0', END)
    text3.delete('1.0', END)
    text4.delete('1.0', END)

#submits the inputted login information and redirects the user to the correct screen given the correct credentails
def submit():
    xpassword = password1.get()
    xusername = username1.get()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM logins")
    logins = mycursor.fetchall()
    found = 0
    if xpassword == "" or xusername == "":
        messagebox.showerror("error", "you must enter a username and password")
    else:
        for i in logins:
            if xusername == i[1] and xpassword == i[2]:
                global ID
                ID = i[0]
                found = 1
        if found == 0:
            messagebox.showerror("error", "Username or password not found")
        elif ID[0] == 'm':
            raise_frame(manager_add)
            update_treeview11()
        elif ID[0] == 't':
            raise_frame(trainer_make_classes)
            update_treeview12()
            update_treeview8()
            update_treeview9()
        elif ID[0] == 'c':
            refresh_alerts()
            raise_frame(client_book)
            refresh_bookings()

#submits and validates a persons information when they are being added to the system
def submit_person():
    rolex = role.get()
    if rolex == 1:
        username_newx = username_new.get()
        password_newx = password_new.get()
        firstnamex = firstname.get()
        lastnamex = lastname.get()
        DOBx = cal1.get_date()
        phonex = phone.get()
        emailx = email.get()
        genderx = gender.get()
        postcodex = postcode.get()
        addressx = address.get()
        qualificationsx = text1.get("1.0",END)
        experiencex = text2.get("1.0",END)
        allright = True
        presence = True
        for i in(firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,username_newx,password_newx):
            if i == "":
                presence = False
        if presence == False:
            messagebox.showerror("error", "please ensure you fill out all fields")
            allright = False
        number = False
        if len(password_newx) < 8:
            allright = False
            messagebox.showerror("error", "your password must be at least 8 characters long")
        for i in firstnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your first name may not contain a number")
            allright = False
        number = False
        for i in lastnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your last name may not contain a number")
            allright = False
        if phonex.isnumeric() == False or len(phonex) != 11:
            messagebox.showerror("error", "your phone number is invalid")
            allright = False
        symbol = False
        for i in emailx:
            if i == '@':
                symbol = True
        space = False
        for i in emailx:
            if i == " ":
                space = True
        if symbol == False or space == True:
            messagebox.showerror("error", "your email address is invalid")
            allright = False
        if genderx == 1:
            genderx = 'male'
        else:
            genderx = 'female'
        if (len(qualificationsx)) == 1:
            qualificationsx = 'none'
        if (len(experiencex)) == 1:
            experiencex = 'none'
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM logins")
        logins = mycursor.fetchall()
        for i in logins:
            if username_newx == i[1]:
                allright = False
                messagebox.showerror("error", "that username has already been taken")
        if allright == True:
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM nextid")
            ids = mycursor.fetchall()
            nextid = ids[-1][0]
            nextid_new = ("m" + str(nextid))
            code = """INSERT INTO logins(ID,username,password_) VALUES (%s,%s,%s)"""
            mycursor.execute(code,(nextid_new,username_newx,password_newx))
            mydb.commit()
            nextid_new2 = int(nextid) + 1
            code = """INSERT INTO nextid(ID,ID2) VALUES (%s,%s)"""
            mycursor.execute(code,(nextid_new2,nextid_new2))
            mydb.commit()
            code = """INSERT INTO managers(ID,firstname,lastname,DOB,phone,email,gender,postcode,address,qualifications,experience) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            mycursor.execute(code,(nextid_new,firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,qualificationsx,experiencex))
            mydb.commit()
            messagebox.showinfo("info", "your user has been succesfully added")
    elif rolex == 2:
        username_newx = username_new.get()
        password_newx = password_new.get()
        firstnamex = firstname.get()
        lastnamex = lastname.get()
        DOBx = cal1.get_date()
        phonex = phone.get()
        emailx = email.get()
        genderx = gender.get()
        postcodex = postcode.get()
        addressx = address.get()
        qualificationsx = text1.get("1.0",END)
        experiencex = text2.get("1.0",END)
        allright = True
        presence = True
        for i in(firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,username_newx,password_newx):
            if i == "":
                presence = False
        if presence == False:
            messagebox.showerror("error", "please ensure you fill out all fields")
            allright = False
        number = False
        if len(password_newx) < 8:
            allright = False
            messagebox.showerror("error", "your password must be at least 8 characters long")
        for i in firstnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your first name may not contain a number")
            allright = False
        number = False
        for i in lastnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your last name may not contain a number")
            allright = False
        if phonex.isnumeric() == False or len(phonex) != 11:
            messagebox.showerror("error", "your phone number is invalid")
            allright = False
        symbol = False
        for i in emailx:
            if i == '@':
                symbol = True
        space = False
        for i in emailx:
            if i == " ":
                space = True
        if symbol == False or space == True:
            messagebox.showerror("error", "your email address is invalid")
            allright = False
        if genderx == 1:
            genderx = 'male'
        else:
            genderx = 'female'
        if (len(qualificationsx)) == 1:
            qualificationsx = 'none'
        if (len(experiencex)) == 1:
            experiencex = 'none'
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM logins")
        logins = mycursor.fetchall()
        for i in logins:
            if username_newx == i[1]:
                allright = False
                messagebox.showerror("error", "that username has already been taken")
        if allright == True:
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM nextid")
            ids = mycursor.fetchall()
            nextid = ids[-1][0]
            nextid_new = ("t" + str(nextid))
            code = """INSERT INTO logins(ID,username,password_) VALUES (%s,%s,%s)"""
            mycursor.execute(code,(nextid_new,username_newx,password_newx))
            mydb.commit()
            nextid_new2 = int(nextid) + 1
            code = """INSERT INTO nextid(ID,ID2) VALUES (%s,%s)"""
            mycursor.execute(code,(nextid_new2,nextid_new2))
            mydb.commit()
            code = """INSERT INTO trainers(ID,firstname,lastname,DOB,phone,email,gender,postcode,address,qualifications,experience) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            mycursor.execute(code,(nextid_new,firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,qualificationsx,experiencex))
            mydb.commit()
            messagebox.showinfo("info", "your user has been succesfully added")
    elif rolex == 3:
        fitnessx = fitness.get()
        smokex = smoke.get()
        heightx = height.get()
        weightx = weight.get()
        username_newx = username_new.get()
        password_newx = password_new.get()
        firstnamex = firstname.get()
        lastnamex = lastname.get()
        DOBx = cal1.get_date()
        phonex = phone.get()
        emailx = email.get()
        genderx = gender.get()
        postcodex = postcode.get()
        addressx = address.get()
        illnessesx = text3.get("1.0",END)
        injuriesx = text4.get("1.0",END)
        allright = True
        presence = True
        for i in(firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,username_newx,password_newx,heightx,weightx):
            if i == "":
                presence = False
        if presence == False:
            messagebox.showerror("error", "please ensure you fill out all fields")
            allright = False
        if heightx.isnumeric() == False:
            messagebox.showerror("error", "your height must be a number")
            allright = False
        if weightx.isnumeric() == False:
            messagebox.showerror("error", "your weight must be a number")
            allright = False
        number = False
        if len(password_newx) < 8:
            allright = False
            messagebox.showerror("error", "your password must be at least 8 characters long")
        for i in firstnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your first name may not contain a number")
            allright = False
        number = False
        for i in lastnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your last name may not contain a number")
            allright = False
        if phonex.isnumeric() == False or len(phonex) != 11:
            messagebox.showerror("error", "your phone number is invalid")
            allright = False
        symbol = False
        for i in emailx:
            if i == '@':
                symbol = True
        space = False
        for i in emailx:
            if i == " ":
                space = True
        if symbol == False or space == True:
            messagebox.showerror("error", "your email address is invalid")
            allright = False
        if smokex == 1:
            smokex = 'no'
        else:
            smokex = 'yes'
        if genderx == 1:
            genderx = 'male'
        else:
            genderx = 'female'
        if (len(illnessesx)) == 1:
            illnessesx = 'none'
        if (len(injuriesx)) == 1:
            injuriesx = 'none'
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM logins")
        logins = mycursor.fetchall()
        for i in logins:
            if username_newx == i[1]:
                allright = False
                messagebox.showerror("error", "that username has already been taken")
        if allright == True:
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM nextid")
            ids = mycursor.fetchall()
            nextid = ids[-1][0]
            nextid_new = ("c" + str(nextid))
            code = """INSERT INTO logins(ID,username,password_) VALUES (%s,%s,%s)"""
            mycursor.execute(code,(nextid_new,username_newx,password_newx))
            mydb.commit()
            nextid_new2 = int(nextid) + 1
            code = """INSERT INTO nextid(ID,ID2) VALUES (%s,%s)"""
            mycursor.execute(code,(nextid_new2,nextid_new2))
            mydb.commit()
            code = """INSERT INTO clients(ID,firstname,lastname,DOB,phone,email,gender,postcode,address,illnesses,injuries,height,weight,smoke,fitness) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            mycursor.execute(code,(nextid_new,firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,illnessesx,injuriesx,heightx,weightx,smokex,fitnessx))
            mydb.commit()
            messagebox.showinfo("info", "your user has been succesfully added")

#resubmits and validates a persons information when their data is being editted
def edit_person():
    mycursor = mydb.cursor()
    selectedID2 = (str(selectedID), )
    code = """DELETE FROM logins WHERE ID=%s LIMIT 1"""
    mycursor.execute(code,selectedID2)
    mydb.commit()
    if selectedID[0] == 'm':
        rolex = 1
    elif selectedID[0] == 't':
        rolex = 2
    else:
        rolex = 3
    if rolex == 1:
        mycursor = mydb.cursor()
        selectedID2 = (str(selectedID), )
        code = """DELETE FROM managers WHERE ID=%s LIMIT 1"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        username_newx = username_new.get()
        password_newx = password_new.get()
        firstnamex = firstname.get()
        lastnamex = lastname.get()
        DOBx = cal1.get_date()
        phonex = phone.get()
        emailx = email.get()
        genderx = gender.get()
        postcodex = postcode.get()
        addressx = address.get()
        qualificationsx = text1.get("1.0",END)
        experiencex = text2.get("1.0",END)
        allright = True
        presence = True
        for i in(firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,username_newx,password_newx):
            if i == "":
                presence = False
        if presence == False:
            messagebox.showerror("error", "please ensure you fill out all fields")
            allright = False
        number = False
        if len(password_newx) < 8:
            allright = False
            messagebox.showerror("error", "your password must be at least 8 characters long")
        for i in firstnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your first name may not contain a number")
            allright = False
        number = False
        for i in lastnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your last name may not contain a number")
            allright = False
        if phonex.isnumeric() == False or len(phonex) != 11:
            messagebox.showerror("error", "your phone number is invalid")
            allright = False
        symbol = False
        for i in emailx:
            if i == '@':
                symbol = True
        space = False
        for i in emailx:
            if i == " ":
                space = True
        if symbol == False or space == True:
            messagebox.showerror("error", "your email address is invalid")
            allright = False
        if genderx == 1:
            genderx = 'male'
        else:
            genderx = 'female'
        if (len(qualificationsx)) == 1:
            qualificationsx = 'none'
        if (len(experiencex)) == 1:
            experiencex = 'none'
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM logins")
        logins = mycursor.fetchall()
        for i in logins:
            if username_newx == i[1]:
                allright = False
                messagebox.showerror("error", "that username has already been taken")
        if allright == True:
            mycursor = mydb.cursor()
            code = """INSERT INTO logins(ID,username,password_) VALUES (%s,%s,%s)"""
            mycursor.execute(code,(selectedID,username_newx,password_newx))
            mydb.commit()
            code = """INSERT INTO managers(ID,firstname,lastname,DOB,phone,email,gender,postcode,address,qualifications,experience) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            mycursor.execute(code,(selectedID,firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,qualificationsx,experiencex))
            mydb.commit()
            messagebox.showinfo("info", "your user has been succesfully edited")
            submit_person_button = HoverButton(manager_add, text = "Submit",height = 2, width = 25,font = (100), command = submit_person, bg = 'black', fg = 'white').place(x=10,y=500)
            clear_all_button = HoverButton(manager_add, text = "Clear All",height = 2, width = 25,font = (100), command = clear_all, bg = 'black', fg = 'white').place(x=300,y=500)
    elif rolex == 2:
        mycursor = mydb.cursor()
        selectedID2 = (str(selectedID), )
        code = """DELETE FROM trainers WHERE ID=%s LIMIT 1"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        username_newx = username_new.get()
        password_newx = password_new.get()
        firstnamex = firstname.get()
        lastnamex = lastname.get()
        DOBx = cal1.get_date()
        phonex = phone.get()
        emailx = email.get()
        genderx = gender.get()
        postcodex = postcode.get()
        addressx = address.get()
        qualificationsx = text1.get("1.0",END)
        experiencex = text2.get("1.0",END)
        allright = True
        presence = True
        for i in(firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,username_newx,password_newx):
            if i == "":
                presence = False
        if presence == False:
            messagebox.showerror("error", "please ensure you fill out all fields")
            allright = False
        number = False
        if len(password_newx) < 8:
            allright = False
            messagebox.showerror("error", "your password must be at least 8 characters long")
        for i in firstnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your first name may not contain a number")
            allright = False
        number = False
        for i in lastnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your last name may not contain a number")
            allright = False
        if phonex.isnumeric() == False or len(phonex) != 11:
            messagebox.showerror("error", "your phone number is invalid")
            allright = False
        symbol = False
        for i in emailx:
            if i == '@':
                symbol = True
        space = False
        for i in emailx:
            if i == " ":
                space = True
        if symbol == False or space == True:
            messagebox.showerror("error", "your email address is invalid")
            allright = False
        if genderx == 1:
            genderx = 'male'
        else:
            genderx = 'female'
        if (len(qualificationsx)) == 1:
            qualificationsx = 'none'
        if (len(experiencex)) == 1:
            experiencex = 'none'
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM logins")
        logins = mycursor.fetchall()
        for i in logins:
            if username_newx == i[1]:
                allright = False
                messagebox.showerror("error", "that username has already been taken")
        if allright == True:
            mycursor = mydb.cursor()
            code = """INSERT INTO logins(ID,username,password_) VALUES (%s,%s,%s)"""
            mycursor.execute(code,(selectedID,username_newx,password_newx))
            mydb.commit()
            code = """INSERT INTO trainers(ID,firstname,lastname,DOB,phone,email,gender,postcode,address,qualifications,experience) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            mycursor.execute(code,(selectedID,firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,qualificationsx,experiencex))
            mydb.commit()
            messagebox.showinfo("info", "your user has been succesfully edited")
            submit_person_button = HoverButton(manager_add, text = "Submit",height = 2, width = 25,font = (100), command = submit_person, bg = 'black', fg = 'white').place(x=10,y=500)
            clear_all_button = HoverButton(manager_add, text = "Clear All",height = 2, width = 25,font = (100), command = clear_all, bg = 'black', fg = 'white').place(x=300,y=500)
    elif rolex == 3:
        mycursor = mydb.cursor()
        selectedID2 = (str(selectedID), )
        code = """DELETE FROM clients WHERE ID=%s LIMIT 1"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        fitnessx = fitness.get()
        smokex = smoke.get()
        heightx = height.get()
        weightx = weight.get()
        username_newx = username_new.get()
        password_newx = password_new.get()
        firstnamex = firstname.get()
        lastnamex = lastname.get()
        DOBx = cal1.get_date()
        phonex = phone.get()
        emailx = email.get()
        genderx = gender.get()
        postcodex = postcode.get()
        addressx = address.get()
        illnessesx = text3.get("1.0",END)
        injuriesx = text4.get("1.0",END)
        allright = True
        presence = True
        for i in(firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,username_newx,password_newx,heightx,weightx):
            if i == "":
                presence = False
        if presence == False:
            messagebox.showerror("error", "please ensure you fill out all fields")
            allright = False
        if heightx.isnumeric() == False:
            messagebox.showerror("error", "your height must be a number")
            allright = False
        if weightx.isnumeric() == False:
            messagebox.showerror("error", "your weight must be a number")
            allright = False
        number = False
        if len(password_newx) < 8:
            allright = False
            messagebox.showerror("error", "your password must be at least 8 characters long")
        for i in firstnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your first name may not contain a number")
            allright = False
        number = False
        for i in lastnamex:
            if i.isnumeric() == True:
                number = True
        if number == True:
            messagebox.showerror("error", "Your last name may not contain a number")
            allright = False
        if phonex.isnumeric() == False or len(phonex) != 11:
            messagebox.showerror("error", "your phone number is invalid")
            allright = False
        symbol = False
        for i in emailx:
            if i == '@':
                symbol = True
        space = False
        for i in emailx:
            if i == " ":
                space = True
        if symbol == False or space == True:
            messagebox.showerror("error", "your email address is invalid")
            allright = False
        if smokex == 1:
            smokex = 'no'
        else:
            smokex = 'yes'
        if genderx == 1:
            genderx = 'male'
        else:
            genderx = 'female'
        if (len(illnessesx)) == 1:
            illnessesx = 'none'
        if (len(injuriesx)) == 1:
            injuriesx = 'none'
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM logins")
        logins = mycursor.fetchall()
        for i in logins:
            if username_newx == i[1]:
                allright = False
                messagebox.showerror("error", "that username has already been taken")
        if allright == True:
            mycursor = mydb.cursor()
            code = """INSERT INTO logins(ID,username,password_) VALUES (%s,%s,%s)"""
            mycursor.execute(code,(selectedID,username_newx,password_newx))
            mydb.commit()
            code = """INSERT INTO clients(ID,firstname,lastname,DOB,phone,email,gender,postcode,address,illnesses,injuries,height,weight,smoke,fitness) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"""
            mycursor.execute(code,(selectedID,firstnamex,lastnamex,DOBx,phonex,emailx,genderx,postcodex,addressx,illnessesx,injuriesx,heightx,weightx,smokex,fitnessx))
            mydb.commit()
            messagebox.showinfo("info", "your user has been succesfully edited")
            submit_person_button = HoverButton(manager_add, text = "Submit",height = 2, width = 25,font = (100), command = submit_person, bg = 'black', fg = 'white').place(x=10,y=500)
            clear_all_button = HoverButton(manager_add, text = "Clear All",height = 2, width = 25,font = (100), command = clear_all, bg = 'black', fg = 'white').place(x=300,y=500)

#deletes a person from the system whilst maintaining referentail integrity
#ensures that if a trainer is deleted all of thier classes and assocaited bookings are also deleted
#ensures that if a client is deleted all of thier associated bookings and fitness information is deleted 
#ensures that all login details are deleted
def delete_person():
    mycursor = mydb.cursor()
    if selectedID[0] == 'c':
        selectedID2 = (str(selectedID), )
        code = """DELETE FROM clients WHERE ID=%s LIMIT 1"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        code = """DELETE FROM bookings WHERE client_ID=%s"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        code = """DELETE FROM tracker WHERE ID=%s"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        code = """DELETE FROM alerts WHERE client_ID=%s"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        code = """DELETE FROM alerts2 WHERE client_ID=%s"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
    elif selectedID[0] == 'm':
        selectedID2 = (str(selectedID), )
        code = """DELETE FROM managers WHERE ID=%s LIMIT 1"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
    elif selectedID[0] == 't':
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM trainers")
        trainers = mycursor.fetchall()
        selectedID2 = (str(selectedID), )
        code = """DELETE FROM trainers WHERE ID=%s LIMIT 1"""
        mycursor.execute(code,selectedID2)
        mydb.commit()
        for i in trainers:
            if i[0] == selectedID:
                name = i[1] + ' ' + i[2]
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM classes")
        classes = mycursor.fetchall()
        for i in classes:
            if i[1] == name:
                cID = i[0]
                print(cID)
                mycursor = mydb.cursor()
                code = """DELETE FROM classes WHERE class_ID = %s AND class_ID = %s"""
                mycursor.execute(code,(cID,cID))
                mydb.commit()
                mycursor = mydb.cursor()
                code = """DELETE FROM bookings WHERE class_ID = %s AND class_ID = %s"""
                mycursor.execute(code,(cID,cID))
                mydb.commit()
        
    selectedID2 = (str(selectedID), )
    code = """DELETE FROM logins WHERE ID=%s LIMIT 1"""
    mycursor.execute(code,selectedID2)
    mydb.commit()
    messagebox.showinfo("info", "your user has been succesfully deleted")
    submit_person_button = HoverButton(manager_add, text = "Submit",height = 2, width = 25,font = (100), command = submit_person, bg = 'black', fg = 'white').place(x=10,y=500)
    clear_all_button = HoverButton(manager_add, text = "Clear All",height = 2, width = 25,font = (100), command = clear_all, bg = 'black', fg = 'white').place(x=300,y=500)

#selects a person from the edit person screen and inputs thier data into the entry fields
def select_person(a):
    global selectedID
    clear_all()
    item = treeview.selection()
    ID = (treeview.item(item)['values'][1])
    raise_frame(manager_add)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM clients")
    clients = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM managers")
    managers = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM trainers")
    trainers = mycursor.fetchall()
    list_ = []
    for i in (clients,managers,trainers):
        for row in i:
            list_.append(row)
    for i in list_:
        if i[0] == ID:
            person = i
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM logins")
    logins = mycursor.fetchall()
    for i in logins:
        if i[0] == ID:
            username = i[1]
            password = i[2]
    selectedID = ID
    if ID[0] == 'm' or ID[0] == 't':
        username_new.set(username)
        password_new.set(password)
        firstname.set(person[1])
        lastname.set(person[2])
        date1 = person[3].split('-')
        cal1.set_date(date(int(date1[0]),int(date1[1]),int(date1[2])))
        phone.set(person[4])
        email.set(person[5])
        if person[6] == 'male':
            gender.set(1)
        else:
            gender.set(2)
        postcode.set(person[7])
        address.set(person[8])
        text1.insert(tk.END,person[9])
        text2.insert(tk.END,person[10])
    else:
        fitness.set(person[14])
        if person[13] == 'no':
            smoke.set(1)
        else:
            smoke.set(2)
        height.set(person[11])
        weight.set(person[12])
        username_new.set(username)
        password_new.set(password)
        firstname.set(person[1])
        lastname.set(person[2])
        date1 = person[3].split('-')
        cal1.set_date(date(int(date1[0]),int(date1[1]),int(date1[2])))
        phone.set(person[4])
        email.set(person[5])
        if person[6] == 'male':
            gender.set(1)
        else:
            gender.set(2)
        postcode.set(person[7])
        address.set(person[8])
        text3.insert(tk.END,person[9])
        text4.insert(tk.END,person[10])
        
    edit_button = HoverButton(manager_add, text = "Submit edited info",height = 2, width = 25,font = (100), command = edit_person, bg = 'black', fg = 'white').place(x=10,y=500)
    delete_button = HoverButton(manager_add, text = "Delete person",height = 2, width = 25,font = (100), command = delete_person, bg = 'black', fg = 'white').place(x=300,y=500)
    

#buffer function that ensures the refine search function runs every key press
def refine_search2(a):
    refine_search()
    refine_search()
    refine_search()

#a function that appropraitly refines the search for users depending on what has been entered into the entry boxes
def refine_search():
    firstnamex2 = firstname_search.get()
    lastnamex2 = lastname_search.get()
    IDx2 = ID_search.get()
    records = treeview.get_children()
    for elements in records:
        treeview.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM clients")
    clients = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM managers")
    managers = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM trainers")
    trainers = mycursor.fetchall()
    list_ = []
    for i in (clients,managers,trainers):
        for row in i:
            list_.append(row)
    list_2 = []
    for i in list_:
        new = (i[1],i[2],i[0])
        list_2.append(new)
    list_ = list_2
    new_list = []
    selectors = []
    if firstnamex2 != "":
        selectors.append(firstnamex2)
    else:
        selectors.append("@")
    if lastnamex2 != "":
        selectors.append(lastnamex2)
    else:
        selectors.append("@")
    if IDx2 != "":
        selectors.append(IDx2)
    else:
        selectors.append("@")
    for i in list_:
        match = True
        if selectors[0] != "@":
            if selectors[0] != i[0][0:len(selectors[0])]:
                match = False
        if selectors[1] != "@":
            if selectors[1] != i[1][0:len(selectors[1])]:
                match = False
        if selectors[2] != "@":
            if selectors[2] != i[2][0:len(selectors[2])]:
                match = False
        if match == True:
            new_list.append(i)
    for row in new_list:
        treeview.insert('','end',text = row[0],values=(row[1],row[2]))

#cancels the search for a user
def cancel_search():
    raise_frame(manager_add)

#brings the client to the search screen    
def search_people():
    raise_frame(search)
    refine_search()

#allows a client to add a new exercise to the fitness tracker
def submit_new_exercise():
    new_exercisex = new_exercise.get()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM tracker")
    tracker = mycursor.fetchall()
    mycursor = mydb.cursor()
    exercises = []
    used = False
    for i in tracker:
        if i[0] == ID:
            exercises.append(eval(i[1]))
    for i in exercises:
        if i[0] == new_exercisex:
            used = True
    number = False
    for i in new_exercisex:
        if i.isnumeric() == True:
            number = True
    if new_exercisex == "":
        messagebox.showerror("error", "please enter an exercise")    
    elif number == True:
        messagebox.showerror("error", "the exercise cannot have a number in it")
    elif used == True:
        messagebox.showerror("error", "you have already used that exercise name")
    else:
        exercise = []
        exercise.append(new_exercisex)
        mycursor = mydb.cursor()
        code = """INSERT INTO tracker(ID,exercise) VALUES (%s,%s)"""
        mycursor.execute(code,(ID,str(exercise)))
        mydb.commit()
        messagebox.showinfo("info", "your exercise has been succesfully added")
        update2()

#selects an exercise from the fitness tracker and places it into a global variable
def select_exercise(x):
    item = treeview2.selection()
    exercise = (treeview2.item(item)['text'])
    global selected_exercise
    selected_exercise = exercise
    
#updates the treeview for the fitness tracker
def update2():
    records2 = treeview2.get_children()
    for elements in records2:
        treeview2.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM tracker")
    tracker = mycursor.fetchall()
    mycursor = mydb.cursor()
    exercises = []
    for i in tracker:
        if i[0] == ID:
            exercises.append(eval(i[1]))
    for i in exercises:
        treeview2.insert('','end',text = i[0],values=(''))

#deletes an exercise that was previously selected for the fitness tracker        
def delete_selected_exercise():
    try:
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM tracker")
        tracker = mycursor.fetchall()
        mycursor = mydb.cursor()
        for i in tracker:
            if i[0] == ID and eval(i[1])[0] == selected_exercise:
                exercise = i[1]
        mycursor = mydb.cursor()
        code = """DELETE FROM tracker WHERE ID = %s AND exercise = %s;"""
        mycursor.execute(code,(ID,exercise))
        mydb.commit()
        update2()
        messagebox.showinfo("info", "your exercise has been succesfully deleted")
    except:
        messagebox.showerror("error", "you must select an exercise")

#adds a data point to the fitness tracker
def add_data():
    weight2x = weight2.get()
    repsx = reps.get()
    setsx = sets.get()
    sesh_timex = sesh_time.get()
    all_right = True
    for i in (weight2x,repsx,setsx,sesh_timex):
        if i == "" or i.isnumeric() == False:
            all_right = False
    if all_right == True:
        current_date = date.today()
        try:
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM tracker")
            tracker = mycursor.fetchall()
            mycursor = mydb.cursor()
            for i in tracker:
                if i[0] == ID and eval(i[1])[0] == selected_exercise:
                    exercise = i[1]
            mycursor = mydb.cursor()
            code = """DELETE FROM tracker WHERE ID = %s AND exercise = %s;"""
            mycursor.execute(code,(ID,exercise))
            mydb.commit()
            new_exercise = eval(exercise)
            new = (str(current_date),weight2x,repsx,setsx,sesh_timex)
            new_exercise.append(new)
            mycursor = mydb.cursor()
            code = """INSERT INTO tracker(ID,exercise) VALUES (%s,%s)"""
            mycursor.execute(code,(ID,str(new_exercise)))
            mydb.commit()
            messagebox.showinfo("info", "your performance has been successfully recorded")
            update2()
        except:
            messagebox.showerror("error", "you must select an exercise")
    else:
        messagebox.showerror("error", "the data you have entered is invalid, please ensure you fill out all fields correctly")

#generates a graph of performance for a specific exercise    
def show_graph_exercise():
    try:
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM tracker")
        tracker = mycursor.fetchall()
        mycursor = mydb.cursor()
        for i in tracker:
            if i[0] == ID and eval(i[1])[0] == selected_exercise:
                exercise = i[1]
        exercise = eval(exercise)
        data = exercise[1:len(exercise)]
        if data == []:
            messagebox.showerror("error", "this exercise has no data to plot")
        else:
            n = len(exercise)-1
            title = exercise[0]
            dates,weight,reps,sets,time = [],[],[],[],[]
            for i in exercise[1:len(exercise)]:
                dates.append(datetime.strptime(i[0], '%Y-%m-%d'))
                weight.append(int(i[1]))
                reps.append(int(i[2]))
                sets.append(int(i[3]))
                time.append(int(i[4]))
            workbook = xlsxwriter.Workbook('tracker.xlsx')
            worksheet = workbook.add_worksheet()
            bold = workbook.add_format({'bold': 1})
            headings = ['Date', 'weight/kg', 'reps/n', 'sets/n', 'time/s']
            worksheet.write_row('A1', headings, bold)
            date_format = workbook.add_format({'num_format': 'dd/mm/yyyy'})
            worksheet.write_column('A2', dates, date_format)
            worksheet.write_column('B2', weight)
            worksheet.write_column('C2', reps)
            worksheet.write_column('D2', sets)
            worksheet.write_column('E2', time)
            chart1 = workbook.add_chart({'type': 'line'})
            chart1.add_series({
                'name':       ['Sheet1', 0, 1],
                'values':     ['Sheet1', 1, 1, n, 1],
                'categories': ['Sheet1', 1, 0, n, 0],
            })
            chart1.add_series({
                'name':       ['Sheet1', 0, 2],
                'categories': ['Sheet1', 1, 0, n, 0],
                'values':     ['Sheet1', 1, 2, n, 2],
            })
            chart1.add_series({
                'name':       ['Sheet1', 0, 3],
                'categories': ['Sheet1', 1, 0, n, 0],
                'values':     ['Sheet1', 1, 3, n, 3],
            })
            chart1.add_series({
                'name':       ['Sheet1', 0, 4],
                'categories': ['Sheet1', 1, 0, n, 0],
                'values':     ['Sheet1', 1, 4, n, 4],
                'y2_axis': 1,
            })
            chart1.set_title ({'name': title})
            chart1.set_x_axis({'name': 'Date'})
            chart1.set_y_axis({'name': 'recorded Data'})
            chart1.set_y2_axis({'name': 'time/s'})
            chart1.set_style(10)
            worksheet.insert_chart('A1', chart1, {'x_scale': 2, 'y_scale': 2})
            workbook.close()
            excel2img.export_img("tracker.xlsx", "tracker chart.png", "Sheet1", 'A1:o29')
            img = Image.open("tracker chart.png")
            img = img.resize((780,500), Image.ANTIALIAS)
            photo =  ImageTk.PhotoImage(img)
            label_t = Label(client_track,image=photo)
            label_t.image = photo
            label_t.place(x=490,y=100)
    except:
        messagebox.showerror("error", "you must select an exercise")

#updates the treeview that displays all classes that a client can book
def update_classes(x):
    my_date = cal2.get_date()
    my_date2 = datetime.strptime(my_date, '%d/%m/%y')
    day = (my_date2.strftime("%a")).lower()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM classes")
    classes = mycursor.fetchall()
    records3 = treeview3.get_children()
    for elements in records3:
        treeview3.delete(elements)
    for i in classes:
        match = False
        days = i[2].split(',')
        for p in days:
            if p == day:
                match = True
        if match == True:
            treeview3.insert('','end',text = i[0],values=(i[1],i[3],i[4],my_date,i[5]))

#books a selected class for a client            
def book_class(x):
    item = treeview3.selection()
    class_ID = (treeview3.item(item)['text'])
    client_ID = ID
    class_date = (treeview3.item(item)['values'][3])
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM bookings")
    bookings = mycursor.fetchall()
    if bookings == []:
        nextID = 1
    else:
        nextID = int(bookings[-1][0])+1
    bookingID = nextID
    mycursor = mydb.cursor()
    code = """INSERT INTO bookings(booking_ID,client_ID,class_ID,class_date) VALUES (%s,%s,%s,%s)"""
    mycursor.execute(code,(bookingID,str(client_ID),class_ID,class_date))
    mydb.commit()
    refresh_bookings()

#deletes a selected booking for a client
def delete_selected_booking(x):
    item = treeview4.selection()
    uID = (treeview4.item(item)['text'])
    mycursor = mydb.cursor()
    code = """DELETE FROM bookings WHERE booking_ID = %s AND booking_ID = %s;"""
    mycursor.execute(code,(str(uID),str(uID)))
    mydb.commit()
    refresh_bookings()

#refreshes the booking treeview for a client and deletes any bookings that are from the past    
def refresh_bookings():
    currentID = ID
    records4 = treeview4.get_children()
    for elements in records4:
        treeview4.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM bookings")
    bookings = mycursor.fetchall()
    for i in range(0,len(bookings)-1):
        mycursor = mydb.cursor()
        code = """DELETE FROM bookings WHERE booking_ID <> %s AND client_ID = %s AND class_ID = %s AND class_date = %s;"""
        mycursor.execute(code,(bookings[i][0],bookings[i][1],bookings[i][2],bookings[i][3]))
        mydb.commit()
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM bookings")
        bookings = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM bookings")
    bookings = mycursor.fetchall()
    for i in bookings:
        d1 = str((date.today())).split('-')
        d2 = i[3].split('/')
        d2.reverse()
        d2 = ["20"+(d2[0]),d2[1],d2[2]]
        if datetime(int(d1[0]),int(d1[1]),int(d1[2])) > datetime(int(d2[0]),int(d2[1]),int(d2[2])):
            delID = i[0]
            mycursor = mydb.cursor()
            code = """DELETE FROM bookings WHERE booking_ID = %s AND booking_ID = %s;"""
            mycursor.execute(code,(str(delID),str(delID)))
            mydb.commit()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM bookings")
    bookings = mycursor.fetchall()
    bookings2 = []
    for i in bookings:
        if i[1] == currentID:
            bookings2.append(i)
    bookings = bookings2
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM classes")
    classes = mycursor.fetchall()
    for i in bookings:
        class_date = i[3]
        newID = i[0]
        for p in classes:
            if i[2] == p[0]:
                start_time = p[3]
                end_time = p[4]
                trainer = p[1]
                overview = p[5]
        treeview4.insert('','end',text = newID,values=(trainer,start_time,end_time,class_date,overview))

#updates the alert buttons for a client
def refresh_alerts():
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM alerts")
    alerts = mycursor.fetchall()
    for i in alerts:
        if i[0] == ID:
            alert_time.set(int(i[2]))
            message_type.set(int(i[1]))

#updates the alert table for a client
def alert_update():
    alert_timex = alert_time.get()
    message_typex = message_type.get()
    currentID = ID
    if currentID[0] == 'c':
        mycursor = mydb.cursor()
        code = """DELETE FROM alerts WHERE client_ID = %s AND client_ID = %s;"""
        mycursor.execute(code,(currentID,currentID))
        mydb.commit()
        mycursor = mydb.cursor()
        code = """INSERT INTO alerts(client_ID,alert_type,alert_time) VALUES (%s,%s,%s);"""
        mycursor.execute(code,(currentID,message_typex,alert_timex))
        mydb.commit()
        refresh_alerts()

#sends an SMS or email or both to a client a certain time before thier class
def send_alert():
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM bookings")
    bookings = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM classes")
    classes = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM alerts")
    alerts = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM clients")
    clients = mycursor.fetchall()
    alert_list = []
    for i in bookings:
        for p in alerts:
            if i[1] == p[0]:
                class_date = i[3]
                alert_type = p[1]
                alert_time = p[2]
                for u in classes:
                    if u[0] == i[2]:
                        class_time = u[3]
                for u in clients:
                    if u[0] == i[1]:
                        email = u[5]
                        phone = u[4]
                
        new = (alert_type,alert_time,class_time,class_date,email,phone)
        alert_list.append(new)
    alert_list2 = []
    for i in alert_list:
        if i[0] == '4' or i[0] == '0' or i[1] == '0':
            pass
        else:
            alert_list2.append(i)
    alert_list = alert_list2
    alert_list2 = []
    d1 = str((date.today())).split('-')
    t1 = (str(datetime.now())[11:-10]).split(':')
    t11 = (int(t1[0])*60) + (int(t1[1]))
    t1 = t11
    for i in alert_list:
        d2 = i[3].split('/')
        d2.reverse()
        d2 = ["20"+(d2[0]),d2[1],d2[2]]
        t2 = i[2].split(':')
        t22 = (int(t2[0])*60) + (int(t2[1])) - (int(i[1]))
        t2 = t22
        if datetime(int(d1[0]),int(d1[1]),int(d1[2])) != datetime(int(d2[0]),int(d2[1]),int(d2[2])):
            pass
        else:
            if t1 != t2:
                pass
            else:
                alert_list2.append(i)
    alert_list = alert_list2
    for i in alert_list:
        if i[0] == '2':
            send_email(i[4],i[1])
        elif i[0] == '1':
            send_text(i[5],i[1])
        elif i[0] == '3':
            print(i[4],i[5])
            send_email(i[4],i[1])
            send_text(i[5],i[1])
    root.after(30000,send_alert)
    
#sends an email to clients
def send_email(receiver,mins):
    port = 465
    smtp_server = "smtp.gmail.com"
    sender_email = "callum.reid2469@gmail.com"
    receiver_email = str(receiver)
    password = "btec12345"
    message = """\
Subject: class alert

your class begins in {} minutes.""".format(str(mins))
    context = ssl.create_default_context()
    with smtplib.SMTP_SSL(smtp_server, port, context=context) as server:
        server.login(sender_email, password)
        server.sendmail(sender_email, receiver_email, message)

#sends a text message to clients
def send_text(receiver,mins):
    receiver = '+44' + str(receiver)
    client = boto3.client("sns",aws_access_key_id="AKIAI6AP5ZTQ6CBWHL5A",aws_secret_access_key="34ToepgjEJ37Qc5oWrhxoEya/dQc7KaksbKXDdUE",region_name="us-east-1")
    client.publish(PhoneNumber=str(receiver),Message="your class begins in {} minutes".format(mins))

#generates a word document report of a clients performance for a selected exercise
def generate_report():
    try:
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM tracker")
        tracker = mycursor.fetchall()
        for i in tracker:
            if i[0] == ID and eval(i[1])[0] == selected_exercise:
                exercise = i[1]
        exercise = eval(exercise)
        data = exercise[1:len(exercise)]
        if data == []:
            messagebox.showerror("error", "this exercise has no data to plot")
        else:
            document = Document()
            sections = document.sections
            for section in sections:
                section.top_margin = Cm(0.5)
                section.bottom_margin = Cm(0.5)
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)
            document.add_heading(selected_exercise, 0)
            for i in data:
                new_term = "on the date {} you did {} sets of {} reps with a weight of {}kg in {} seconds".format(i[0],i[3],i[2],i[1],i[4])
                p = document.add_paragraph(new_term)
            show_graph_exercise()
            img = Image.open("tracker chart.png")
            img = img.resize((550,350), Image.ANTIALIAS)
            img.save("tracker chart.png")
            p = document.add_picture("tracker chart.png")
            path = easygui.diropenbox()
            document.save('{}\{}.docx'.format(path,selected_exercise))
            messagebox.showinfo("info", "your report has been succesfully generated")
    except:
        messagebox.showerror("error", "you must select an exercise")

#submits entered information for a trainer to make a class
def create_class():
    mon = monday.get()
    tue = tuesday.get()
    wed = wednesday.get()
    thu = thursday.get()
    fri = friday.get()
    sat = saturday.get()
    sun = sunday.get()
    time1x = time1.get()
    time2x = time2.get()
    class_durationx = class_duration.get()
    overview = text984.get("1.0",END)
    days = []
    if mon == 1:
        days.append("mon")
    if tue == 1:
        days.append("tue")
    if wed == 1:
        days.append("wed")
    if thu == 1:
        days.append("thu")
    if fri == 1:
        days.append("fri")
    if sat == 1:
        days.append("sat")
    if sun == 1:
        days.append("sun")
    verified = True
    if days == []:
        verified = False
    if class_durationx == "":
        verified = False
    if len(overview) == 1:
        verified = False
    try:
        duration = int(class_durationx)
        if verified == True:
            start_time = (str(time1x) + ":" + str(time2x))
            end_time_mins = (((int(time1x) * 60) + int(time2x)) + duration)
            temp = str(end_time_mins % 60)
            if len(temp) == 1:
                temp = "0" + temp
            end_time = str(end_time_mins // 60) + ":" + temp
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM trainers")
            trainers = mycursor.fetchall()
            for i in trainers:
                if i[0] == ID:
                    trainer = (str(i[1]) + " " + str(i[2]))
            temp = start_time.split(":")
            if len(temp[1]) == 1:
                start_time = temp[0] + ":" + "0" + temp[1]
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM classes")
            classes = mycursor.fetchall()
            if classes == []:
                nextID = 1
            else:
                nextID = int(classes[-1][0])+1
            classID = nextID
            days = str(days)
            days = days.replace("'","")
            days = days.replace("[","")
            days = days.replace("]","")
            days = days.replace(" ","")
            mycursor = mydb.cursor()
            code = """INSERT INTO classes(class_ID,trainer,days,start_time,end_time,overview) VALUES (%s,%s,%s,%s,%s,%s);"""
            mycursor.execute(code,(classID,trainer,days,start_time,end_time,overview))
            mydb.commit()
            update_treeview8()
            messagebox.showinfo("info", "your class has been created")
        else:
            messagebox.showerror("error", "you must fill in all fields")
    except:
        messagebox.showerror("error", "class duration must be a number")

#allows a trainer to delete a previously made class
def select_class(x):
    item = treeview8.selection()
    cID = (treeview8.item(item)['text'])
    MsgBox = tk.messagebox.askquestion ('Delete','Are you sure you want to delete this class',icon = 'warning')
    if MsgBox == 'yes':
        mycursor = mydb.cursor()
        code = """DELETE FROM classes WHERE class_ID = %s AND class_ID = %s"""
        mycursor.execute(code,(cID,cID))
        mydb.commit()
        mycursor = mydb.cursor()
        code = """DELETE FROM bookings WHERE class_ID = %s AND class_ID = %s"""
        mycursor.execute(code,(cID,cID))
        mydb.commit()
        update_treeview8()
        
#updates the treeview that contains all of a trainers previously made classes
def update_treeview8():
    global ID
    records8 = treeview8.get_children()
    for elements in records8:
        treeview8.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM trainers")
    trainers = mycursor.fetchall()
    for i in trainers:
        if i[0] == ID:
            name = str(i[1] + " " + i[2])
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM classes")
    classes = mycursor.fetchall()
    for i in classes:
        if i[1] == name:
            treeview8.insert('','end',text = i[0],values=(i[2],i[3],i[4],i[5]))

#updates a treeview with the next 30 days of classes so that a trainer can viewwho is booked into them
def update_treeview9():
    global ID
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM trainers")
    trainers = mycursor.fetchall()
    for i in trainers:
        if i[0] == ID:
            Tname = str(i[1]) + " " + str(i[2])
    records9 = treeview9.get_children()
    for elements in records9:
        treeview9.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM classes")
    classes = mycursor.fetchall()
    dates = []
    d = datetime.today() + timedelta(days=30)
    for i in range(0,31):
        dates.append(d)
        d = d - timedelta(days=1)
    dates2 = []
    for i in dates:
        date = str(i)[0:10]
        day = (datetime.strptime(date.replace('-',' '), '%Y %m %d').strftime('%a')).lower()
        new = (date,day)
        dates2.append(new)
    dates2.reverse()
    for i in dates2:
        for p in classes:
            confirm = False
            if str(p[1]) == Tname:
                confirm = True
            found = False
            for q in p[2].split(','):
                if q == i[1]:
                    found = True
            if found == True and confirm == True:
                treeview9.insert('','end',text = i[0],values=(i[0],p[0],p[3],p[4],p[5]))
            else:
                pass
                                
#allows a trainer to view who has booked into a selected class
def check_bookings(x):
    global ID
    records10 = treeview10.get_children()
    for elements in records10:
        treeview10.delete(elements)
    item = treeview9.selection()
    date = (treeview9.item(item)['text'])
    date = date.split('-')
    date.reverse()
    date = (date[0],date[1],date[2][2:])
    date = (date[0]+'/'+date[1]+'/'+date[2])
    classID = (treeview9.item(item)['values'][1])
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM bookings")
    bookings = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM clients")
    clients = mycursor.fetchall()
    for i in bookings:
        if (str(i[2]) == str(classID)) and (str(i[3]) == str(date)):
            clientID = i[1]
            for i in clients:
                if i[0] == clientID:
                    name = (i[1] + " " + i[2])
                    phone = i[4]
            treeview10.insert('','end',text = i[0],values=(clientID,classID,date,name,phone))

#allows a trainer to mark someone as present at a class    
def mark_as_present(x):
    d = str(datetime.today())[0:10]
    date = d.split('-')
    date.reverse()
    date = (date[0],date[1],date[2][2:])
    datet = (date[0]+'/'+date[1]+'/'+date[2])
    item = treeview10.selection()
    clientID = (treeview10.item(item)['values'][0])
    classID = (treeview10.item(item)['values'][1])
    date = (treeview10.item(item)['values'][2])
    name = (treeview10.item(item)['values'][3])
    phone = (treeview10.item(item)['values'][4])
    if date == datet:
        MsgBox = tk.messagebox.askquestion ('record','Do you wish to record this person as present',icon = 'info')
        if MsgBox == 'yes':
            mycursor = mydb.cursor()
            code = """INSERT INTO attendances(clientID,classID,date_,name_,phone) VALUES (%s,%s,%s,%s,%s);"""
            mycursor.execute(code,(clientID,classID,date,name,phone))
            mydb.commit()
            messagebox.showinfo("info", "this attendance has been rcorded")

            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM attendances")
            attendances = mycursor.fetchall()
            for i in attendances:
                client_ID = i[0]
                class_ID = i[1]
                date = i[2]
                name = i[3]
                phone = i[4]
                code = """DELETE FROM attendances WHERE clientID = %s AND classID = %s AND date_ = %s AND name_ = %s AND phone = %s """
                mycursor.execute(code,(client_ID,class_ID,date,name,phone))
                mydb.commit()
                code = """INSERT INTO attendances(clientID,classID,date_,name_,phone) VALUES (%s,%s,%s,%s,%s);"""
                mycursor.execute(code,(client_ID,class_ID,date,name,phone))
                mydb.commit()

    else:
        messagebox.showerror("error", "you cannot record attendance for a class that is not today")

#allows a manager to delete a message on the message board
def delete_message(x):
    global ID
    item = treeview11.selection()
    IDx = (treeview11.item(item)['values'][0])
    name = (treeview11.item(item)['values'][1])
    message = (treeview11.item(item)['values'][2])
    if ID != IDx:
        messagebox.showinfo("message", message)
    else:
        MsgBox = tk.messagebox.askquestion ('delete message','Do you wish to delete this message',icon = 'info')
        if MsgBox == 'yes':
            code = """DELETE FROM messages WHERE sender_ID = %s AND sender_name = %s AND message = %s"""
            mycursor.execute(code,(IDx,name,message))
            mydb.commit()
            update_treeview11()

#allows a manager to send a message to the message board
def send_message():
    global ID
    message = text_message.get("1.0",END)
    if int(len(message)) < int(2):
        messagebox.showerror("error", "you must input a message")
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM managers")
    managers = mycursor.fetchall()
    for i in managers:
        if i[0] == ID:
            name = i[1] + ' ' + i[2]
    code = """INSERT INTO messages(sender_ID,sender_name,message) VALUES (%s,%s,%s);"""
    message = message[0:len(message)-1] + ' - sent on ' + (str(datetime.today())[:-7])
    mycursor.execute(code,(ID,name,message))
    mydb.commit()
    update_treeview11()
    messagebox.showinfo("info", "this message has been sent")
    inform_clients(name,message)

#updates the treeview for the message board
def update_treeview11():
    global ID
    records11 = treeview11.get_children()
    for elements in records11:
        treeview11.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM messages")
    messages = mycursor.fetchall()
    for i in messages:
        treeview11.insert('','end',text = i[0],values=(i[0],i[1],i[2]))

#allows a manager to delete a message on the message board
def delete_message2(x):
    global ID
    item = treeview12.selection()
    IDx = (treeview12.item(item)['values'][0])
    name = (treeview12.item(item)['values'][1])
    message = (treeview12.item(item)['values'][2])
    if ID != IDx:
        messagebox.showinfo("message", message)
    else:
        MsgBox = tk.messagebox.askquestion ('delete message','Do you wish to delete this message',icon = 'info')
        if MsgBox == 'yes':
            code = """DELETE FROM messages WHERE sender_ID = %s AND sender_name = %s AND message = %s"""
            mycursor.execute(code,(IDx,name,message))
            mydb.commit()
            update_treeview12()

#allows a manager to send a message to the message board
def send_message2():
    global ID
    message = text_message2.get("1.0",END)
    if int(len(message)) < int(2):
        messagebox.showerror("error", "you must input a message")
    else:
        mycursor = mydb.cursor()
        mycursor.execute("SELECT * FROM trainers")
        trainers = mycursor.fetchall()
        for i in trainers:
            if i[0] == ID:
                name = i[1] + ' ' + i[2]
        message = message[0:len(message)-1] + ' - sent on ' + (str(datetime.today())[:-7])
        code = """INSERT INTO messages(sender_ID,sender_name,message) VALUES (%s,%s,%s);"""
        mycursor.execute(code,(ID,name,message))
        mydb.commit()
        update_treeview12()
        messagebox.showinfo("info", "this message has been sent")
        inform_clients(name,message)
    
#updates the treeview for the message board
def update_treeview12():
    global ID
    records12 = treeview12.get_children()
    for elements in records12:
        treeview12.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM messages")
    messages = mycursor.fetchall()
    for i in messages:
        treeview12.insert('','end',text = i[0],values=(i[0],i[1],i[2]))

#updates the alert buttons for the client    
def refresh_alerts2():
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM alerts2")
    alerts = mycursor.fetchall()
    for i in alerts:
        if i[0] == ID:
            message_alert_type.set(int(i[1]))

#updates the alert table for the client
def alert_update2():
    message_typex = message_alert_type.get()
    currentID = ID
    if currentID[0] == 'c':
        mycursor = mydb.cursor()
        code = """DELETE FROM alerts2 WHERE client_ID = %s AND client_ID = %s;"""
        mycursor.execute(code,(currentID,currentID))
        mydb.commit()
        mycursor = mydb.cursor()
        code = """INSERT INTO alerts2(client_ID,alert_type) VALUES (%s,%s);"""
        mycursor.execute(code,(currentID,message_typex))
        mydb.commit()
        refresh_alerts()

#updates the treeview for the message board
def update_treeview13():
    global ID
    records13 = treeview13.get_children()
    for elements in records13:
        treeview13.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM messages")
    messages = mycursor.fetchall()
    for i in messages:
        treeview13.insert('','end',text = i[0],values=(i[0],i[1],i[2]))
    
#sends an alert to client that a message has been sent to the message board
def inform_clients(person,message):
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM alerts2")
    alerts = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM clients")
    clients = mycursor.fetchall()
    for i in alerts:
        for p in clients:
            if i[0] == p[0]:
                email = p[5]
                phone = p[4]
        if i[1] == '4' or i[0] == '0' or i[1] == '0':
            pass
        elif i[1] == '1':
            receiver = '+44' + str(phone)
            client = boto3.client("sns",aws_access_key_id="AKIAI6AP5ZTQ6CBWHL5A",aws_secret_access_key="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",region_name="us-east-1")
            client.publish(PhoneNumber=str(receiver),Message="a message has been sent by {}: {}".format(person,message))
        elif i[1] == '2':
            port = 465
            smtp_server = "smtp.gmail.com"
            sender_email = "callum.reid2469@gmail.com"
            receiver_email = email
            password = "btec12345"
            message2 = """\
Subject: message alert

a message has been sent by {}: {}.""".format(person,message)
            context = ssl.create_default_context()
            with smtplib.SMTP_SSL(smtp_server, port, context=context) as server:
                server.login(sender_email, password)
                server.sendmail(sender_email, receiver_email, message2)
        else:
            port = 465
            smtp_server = "smtp.gmail.com"
            sender_email = "callum.reid2469@gmail.com"
            receiver_email = email
            password = "btec12345"
            message2 = """\
Subject: message alert

a message has been sent by {}: {}.""".format(person,message)
            context = ssl.create_default_context()
            with smtplib.SMTP_SSL(smtp_server, port, context=context) as server:
                server.login(sender_email, password)
                server.sendmail(sender_email, receiver_email, message2)
            receiver = '+44' + str(phone)
            client = boto3.client("sns",aws_access_key_id="AKIAI6AP5ZTQ6CBWHL5A",aws_secret_access_key="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",region_name="us-east-1")
            client.publish(PhoneNumber=str(receiver),Message="a message has been sent by {}: {}".format(person,message))
            
#allows someone to view a message
def view_message(x):
    item = treeview13.selection()
    message = (treeview13.item(item)['values'][2])
    messagebox.showinfo("message", message)

#generates a graph for the manager based on a specific selection of data
def show_graph2():
    item = treeview14.selection()
    report = (treeview14.item(item)['values'])
    report2 = ''
    for i in report:
        report2 = str(report2) + ' ' + str(i)
    report = report2[1:]
    if report == '':
        messagebox.showerror("error", "no option selected")
    else:
        if report == 'trainer attendances':
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM trainers")
            trainers = mycursor.fetchall()
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM attendances")
            attendances = mycursor.fetchall()
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM classes")
            classes = mycursor.fetchall()
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM bookings")
            bookings = mycursor.fetchall()
            names = []
            for i in trainers:
                names.append(i[1]+' '+i[2])
            attendance_count = []
            for i in names:
                attendance = 0
                for q in attendances:
                    classID = q[1]
                    for p in classes:
                        if classID == p[0] and i == p[1]:
                            attendance = attendance + 1
                attendance_count.append(attendance)
            booking_count = []
            for i in names:
                booking = 0
                for q in bookings:
                    classID = q[2]
                    for p in classes:
                        if classID == p[0] and i == p[1]:
                            booking = booking + 1
                booking_count.append(booking)
            workbook = xlsxwriter.Workbook('report1.xlsx')
            worksheet = workbook.add_worksheet()
            bold = workbook.add_format({'bold': 1})
            headings = ['name','booking count','attendance count']
            worksheet.write_row('A1', headings, bold)
            worksheet.write_column('A2', names)
            worksheet.write_column('B2', booking_count)
            worksheet.write_column('C2', attendance_count)
            n = len(names)
            chart1 = workbook.add_chart({'type': 'bar'})
            chart1.add_series({
                'name':       ['Sheet1', 0, 1],
                'values':     ['Sheet1', 1, 1, n, 1],
                'categories': ['Sheet1', 1, 0, n, 0],
            })
            chart1.add_series({
                'name':       ['Sheet1', 0, 2],
                'values':     ['Sheet1', 1, 2, n, 2],
                'categories': ['Sheet1', 1, 0, n, 0],
            })
            chart1.set_title ({'name': 'trainer attendances'})
            chart1.set_x_axis({'name': 'number of people'})
            chart1.set_y_axis({'name': 'names'})
            chart1.set_style(10)
            worksheet.insert_chart('A1', chart1, {'x_scale': 2, 'y_scale': 2})
            workbook.close()
            excel2img.export_img("report1.xlsx", "report chart.png", "Sheet1", 'A1:o29')
            img = Image.open("report chart.png")
            img = img.resize((780,500), Image.ANTIALIAS)
            photo =  ImageTk.PhotoImage(img)
            label_t = Label(manager_report,image=photo)
            label_t.image = photo
            label_t.place(x=420,y=100)
        elif report == 'attendances vs times':
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM classes")
            classes = mycursor.fetchall()
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM attendances")
            attendances = mycursor.fetchall()
            times = ['01:00','02:00','03:00','04:00','05:00','06:00','07:00','08:00','09:00','10:00','11:00','12:00',
                     '13:00','14:00','15:00','16:00','17:00','18:00','19:00','20:00','21:00','22:00','23:00','24:00']
            data = []
            for i in times:
                temp = i[0:2]
                counter = 0
                for p in attendances:
                    for q in classes:
                        if p[1] == q[0]:
                            time = q[3]
                            if len(time) < 5:
                                time = '0' + str(time)
                    if temp == time[0:2]:
                        counter = counter + 1
                data.append(counter)
            workbook = xlsxwriter.Workbook('report1.xlsx')
            worksheet = workbook.add_worksheet()
            bold = workbook.add_format({'bold': 1})
            headings = ['time','attendances']
            worksheet.write_row('A1', headings, bold)
            worksheet.write_column('A2', times)
            worksheet.write_column('B2', data)
            n = len(times)
            chart1 = workbook.add_chart({'type': 'bar'})
            chart1.add_series({
                'name':       ['Sheet1', 0, 1],
                'values':     ['Sheet1', 1, 1, n, 1],
                'categories': ['Sheet1', 1, 0, n, 0],
            })
            chart1.set_title ({'name': 'attendances vs time of day'})
            chart1.set_x_axis({'name': 'number of people'})
            chart1.set_y_axis({'name': 'times'})
            chart1.set_style(10)
            worksheet.insert_chart('A1', chart1, {'x_scale': 2, 'y_scale': 2})
            workbook.close()
            excel2img.export_img("report1.xlsx", "report chart.png", "Sheet1", 'A1:o29')
            img = Image.open("report chart.png")
            img = img.resize((780,500), Image.ANTIALIAS)
            photo =  ImageTk.PhotoImage(img)
            label_t = Label(manager_report,image=photo)
            label_t.image = photo
            label_t.place(x=420,y=100)
        elif report == 'attendances (last 30 days)':
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM attendances")
            attendances = mycursor.fetchall()
            dates = []
            d = datetime.today()
            for i in range(0,30):
                dates.append(str(d)[2:10])
                d = d - timedelta(days=1)
            attendances_dates = []
            for i in attendances:
                    date = (i[2].split('/'))
                    date.reverse()
                    d2 = ''
                    for q in date:
                        d2 = d2 + q + '-'
                    d2 = d2[0:-1]
                    attendances_dates.append(d2)
            data = []
            for i in dates:
                counter = 0
                for q in attendances_dates:
                    if i == q:
                        counter = counter + 1
                data.append(counter)
            workbook = xlsxwriter.Workbook('report1.xlsx')
            worksheet = workbook.add_worksheet()
            bold = workbook.add_format({'bold': 1})
            headings = ['dates','attendances']
            worksheet.write_row('A1', headings, bold)
            worksheet.write_column('A2', dates)
            worksheet.write_column('B2', data)
            n = len(data)
            chart1 = workbook.add_chart({'type': 'area'})
            chart1.add_series({
                'name':       ['Sheet1', 0, 1],
                'values':     ['Sheet1', 1, 1, n, 1],
                'categories': ['Sheet1', 1, 0, n, 0],
                'trendline': {
                'type': 'polynomial',
                'name': 'trend',
                'order': 2,
                'forward': 0.5,
                'backward': 0.5,
                'line': {
                    'color': 'red',
                    'width': 1,
                    'dash_type': 'long_dash',
                    }},
            })
            chart1.set_title ({'name': 'attendances for the last 30 days'})
            chart1.set_x_axis({'name': 'date'})
            chart1.set_y_axis({'name': 'number of people'})
            chart1.set_style(10)
            worksheet.insert_chart('A1', chart1, {'x_scale': 2, 'y_scale': 2})
            workbook.close()
            excel2img.export_img("report1.xlsx", "report chart.png", "Sheet1", 'A1:o29')
            img = Image.open("report chart.png")
            img = img.resize((780,500), Image.ANTIALIAS)
            photo =  ImageTk.PhotoImage(img)
            label_t = Label(manager_report,image=photo)
            label_t.image = photo
            label_t.place(x=420,y=100)
        elif report == 'attendances vs class types':
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM attendances")
            attendances = mycursor.fetchall()
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM classes")
            classes = mycursor.fetchall()
            types = []
            for i in classes:
                dupe = False
                for p in types:
                    if i[5].replace('\n','') == p:
                        dupe = True
                if dupe == False:
                    types.append(i[5].replace('\n',''))
            data = []
            for i in types:
                counter = 0
                for q in attendances:
                    for p in classes:
                        if p[0] == q[1] and i == p[5].replace('\n',''):
                            counter = counter + 1
                data.append(counter)
            workbook = xlsxwriter.Workbook('report1.xlsx')
            worksheet = workbook.add_worksheet()
            bold = workbook.add_format({'bold': 1})
            headings = ['types','attendances']
            worksheet.write_row('A1', headings, bold)
            worksheet.write_column('A2', types)
            worksheet.write_column('B2', data)
            n = len(data)
            chart1 = workbook.add_chart({'type': 'doughnut'})
            chart1.add_series({
                'name':       ['Sheet1', 0, 1],
                'values':     ['Sheet1', 1, 1, n, 1],
                'categories': ['Sheet1', 1, 0, n, 0],
                'data_labels': {'percentage': True},
            })
            chart1.set_title ({'name': 'attendances vs class types'})
            chart1.set_style(10)
            worksheet.insert_chart('A1', chart1, {'x_scale': 2, 'y_scale': 2})
            workbook.close()
            excel2img.export_img("report1.xlsx", "report chart.png", "Sheet1", 'A1:o29')
            img = Image.open("report chart.png")
            img = img.resize((780,500), Image.ANTIALIAS)
            photo =  ImageTk.PhotoImage(img)
            label_t = Label(manager_report,image=photo)
            label_t.image = photo
            label_t.place(x=420,y=100)
        else:
            temp = report.split(' ')
            cID = temp[-1]
            temp = temp[2:-3]
            selected_exercise = ''
            for i in temp:
                selected_exercise = selected_exercise + ' ' +i
            selected_exercise = selected_exercise[1:]
            mycursor = mydb.cursor()
            mycursor.execute("SELECT * FROM tracker")
            tracker = mycursor.fetchall()
            mycursor = mydb.cursor()
            for i in tracker:
                if str(i[0]) == str(cID) and str(eval(i[1])[0]) == str(selected_exercise):
                    exercise = i[1]
            exercise = eval(exercise)
            data = exercise[1:len(exercise)]
            if data == []:
                messagebox.showerror("error", "this exercise has no data to plot")
            else:
                n = len(exercise)-1
                title = exercise[0]
                dates,weight,reps,sets,time = [],[],[],[],[]
                for i in exercise[1:len(exercise)]:
                    dates.append(datetime.strptime(i[0], '%Y-%m-%d'))
                    weight.append(int(i[1]))
                    reps.append(int(i[2]))
                    sets.append(int(i[3]))
                    time.append(int(i[4]))
                workbook = xlsxwriter.Workbook('tracker.xlsx')
                worksheet = workbook.add_worksheet()
                bold = workbook.add_format({'bold': 1})
                headings = ['Date', 'weight/kg', 'reps/n', 'sets/n', 'time/s']
                worksheet.write_row('A1', headings, bold)
                date_format = workbook.add_format({'num_format': 'dd/mm/yyyy'})
                worksheet.write_column('A2', dates, date_format)
                worksheet.write_column('B2', weight)
                worksheet.write_column('C2', reps)
                worksheet.write_column('D2', sets)
                worksheet.write_column('E2', time)
                chart1 = workbook.add_chart({'type': 'line'})
                chart1.add_series({
                    'name':       ['Sheet1', 0, 1],
                    'values':     ['Sheet1', 1, 1, n, 1],
                    'categories': ['Sheet1', 1, 0, n, 0],
                })
                chart1.add_series({
                    'name':       ['Sheet1', 0, 2],
                    'categories': ['Sheet1', 1, 0, n, 0],
                    'values':     ['Sheet1', 1, 2, n, 2],
                })
                chart1.add_series({
                    'name':       ['Sheet1', 0, 3],
                    'categories': ['Sheet1', 1, 0, n, 0],
                    'values':     ['Sheet1', 1, 3, n, 3],
                })
                chart1.add_series({
                    'name':       ['Sheet1', 0, 4],
                    'categories': ['Sheet1', 1, 0, n, 0],
                    'values':     ['Sheet1', 1, 4, n, 4],
                    'y2_axis': 1,
                })
                chart1.set_title ({'name': title})
                chart1.set_x_axis({'name': 'Date'})
                chart1.set_y_axis({'name': 'recorded Data'})
                chart1.set_y2_axis({'name': 'time/s'})
                chart1.set_style(10)
                worksheet.insert_chart('A1', chart1, {'x_scale': 2, 'y_scale': 2})
                workbook.close()
                excel2img.export_img("tracker.xlsx", "tracker chart.png", "Sheet1", 'A1:o29')
                img = Image.open("tracker chart.png")
                img = img.resize((780,500), Image.ANTIALIAS)
                photo =  ImageTk.PhotoImage(img)
                label_t = Label(manager_report,image=photo)
                label_t.image = photo
                label_t.place(x=420,y=100)
            
            
#generates a report for the manager based on a specific set of data
def gen_report2():
    path = easygui.diropenbox()
    selected_exercise = 'report'
    show_graph2()
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    document.add_heading(selected_exercise, 0)
    img = Image.open("report chart.png")
    img = img.resize((550,350), Image.ANTIALIAS)
    img.save("report chart.png")
    p = document.add_picture("report chart.png")
    path = easygui.diropenbox()
    document.save('{}\{}.docx'.format(path,selected_exercise))
    messagebox.showinfo("info", "your report has been succesfully generated")

#fills the database with random data (purely used for testing)
def DB_fill():
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM clients")
    clients = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM classes")
    classes = mycursor.fetchall()
    code = """DELETE FROM attendances;"""
    mycursor.execute(code)
    mydb.commit()
    dates = []
    d = datetime.today()
    for i in range(0,40):
        dates.append(str(d)[2:10])
        d = d - timedelta(days=1)
    dates2 = []
    for i in dates: 
        date = i.split('-')
        date.reverse()
        d2 = ''
        for q in date:
            d2 = d2 + q + '/'
        d2 = d2[0:-1]
        dates2.append(d2)
    for i in dates2:
        for q in range(0,random.randint(15,20)):
            ran = clients[random.randint(0,len(clients)-1)]
            cID = ran[0]
            cname = ran[1] + ' ' + ran[2]
            cphone = ran[4]
            ran = classes[random.randint(0,len(classes)-1)]
            classID = ran[0]
            date = i
            code = """INSERT INTO attendances(clientID,classID,date_,name_,phone) VALUES (%s,%s,%s,%s,%s);"""
            mycursor.execute(code,(cID,classID,date,cname,cphone))
            mydb.commit()

#updates the graph generation treeview with all possible graphs or reports
def update_treeview14():
    records14 = treeview14.get_children()
    for elements in records14:
        treeview14.delete(elements)
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM clients")
    clients = mycursor.fetchall()
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM tracker")
    tracker = mycursor.fetchall()
    exercises = []
    for i in tracker:
        cID = i[0]
        exercise = eval(i[1])[0]
        for q in clients:
            if cID == q[0]:
                name = q[1] + ' ' + q[2]
        quote = name + ' ' + exercise +' '+ 'ID = '+cID
        exercises.append(quote)
    reports2 = []
    reports = ['trainer attendances','attendances vs times','attendances (last 30 days)','attendances vs class types']
    for i in reports:
        reports2.append(i)
    for i in exercises:
        reports2.append(i)
    for i in reports2:
        treeview14.insert('','end',text = i,values=(i))

#brings the user to an internet map showing all locations of bayburn fitness
def show_map():
    m = folium.Map(location=[54.672978 , -5.7444289],zoom_start=12,tiles='Stamen Terrain')
    folium.Marker([54.672978 , -5.7444289], popup='<i>10 Fort Rd, Helens Bay, Bangor BT19 1LA</i>', tooltip='Helens Bay').add_to(m)
    folium.Marker([54.645700,-5.826770], popup='<i>Holywood BT18 0LL</i>', tooltip='Seapark').add_to(m)
    folium.Marker([54.673293,-5.737265], popup='<i>22 Grey Point, Helens Bay, Bangor BT19 1LE</i>', tooltip='Helens Bay Beach').add_to(m)
    m.save('index.html')
    webbrowser.open('file://' + os.path.realpath('index.html'))
   
#6:35                
################################### screen setup ###############################################
root = Tk()
root.geometry("1920x1080")
root.state("zoomed")
root.title("Bayburn Fitness")

login = Frame(root,width=1920, height=1080)

manager_add = Frame(root,width=1920, height=1080)
search = Frame(root,width=300, height=300)
manager_info = Frame(root,width=1920, height=1080)
manager_report = Frame(root,width=1920, height=1080)

trainer_make_classes = Frame(root,width=1920, height=1080)
trainer_info = Frame(root,width=1920, height=1080)
trainer_see_bookings = Frame(root,width=1920, height=1080)

client_book = Frame(root,width=1920, height=1080)
client_info = Frame(root,width=1920, height=1080)
client_track = Frame(root,width=1920, height=1080)

for frame in(login,manager_add,manager_info,manager_report,trainer_make_classes,trainer_info,trainer_see_bookings,client_book,client_info,client_track,search):
    frame.grid(row=0,column = 0, sticky= 'news')


#################################### login ###############################################

    
username1 = StringVar()
password1 = StringVar()
load = Image.open("login background.jpg")
render = ImageTk.PhotoImage(load)
img = Label(login, image=render)
img.image = render
img.place(x=0, y=0)

canvas3 = Canvas(login,bg = 'light grey', width=350, height=200).place(x=450,y=280)
canvas4 = Canvas(login,bg = 'light grey', width=350, height=100).place(x=450,y=100)
labe29 = Label(login, text = "Login", font = ('bold',46)).place(x=550,y=115)
username_entry = Entry(login, textvariable = username1, font = (42)).place(x=570,y=300)
password_entry = Entry(login, textvariable = password1, show="*", font = (42)).place(x=570,y=350)
label2 = Label(login, text = "Username:",font = (46)).place(x=460,y=300)
label3 = Label(login, text = "Password:",font = (46)).place(x=460,y=350)
submit_button = HoverButton(login, text = "Login", command = submit, bg = 'black', fg = 'white', height = 4, width = 40).place(x=490,y=400)



################################ manager_add ############################################
canvas1 = Canvas(manager_add,bg = 'light grey', width=700, height=170).place(x=600,y=120)
canvas2 = Canvas(manager_add,bg = 'light grey', width=700, height=350).place(x=600,y=310)
manager_add_button1 = HoverButton(manager_add, text = "add people", command = manager_add_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
manager_info_button1 = HoverButton(manager_add, text = "information board", command = manager_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
manager_report_button1 = HoverButton(manager_add, text = "data reports", command = manager_report_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button1 = HoverButton(manager_add, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)
label5 = Label(manager_add, text = "Add a person", font = ('bold',36)).place(x=500,y=45)
label13 = Label(manager_add, text = "First Name:", font = (36)).place(x=10,y=120)
firstname = StringVar()
entry1 = Entry(manager_add, textvariable = firstname, font = (36)).place(x=160,y=120)
label14 = Label(manager_add, text = "Last Name:", font = (36)).place(x=10,y=150)
lastname = StringVar()
entry2 = Entry(manager_add, textvariable = lastname, font = (36)).place(x=160,y=150)
label18 = Label(manager_add, text = "Date Of Birth:", font = (36)).place(x=10,y=180)
cal1 = DateEntry(manager_add,width=30,bg="darkblue",fg="white",maxdate = date.today())
cal1.place(x=160,y=180)
label16 = Label(manager_add, text = "Phone Number:", font = (36)).place(x=10,y=210)
phone = StringVar()
entry4 = Entry(manager_add, textvariable = phone, font = (36)).place(x=160,y=210)
label16 = Label(manager_add, text = "Email Address:", font = (36)).place(x=10,y=240)
email = StringVar()
entry5 = Entry(manager_add, textvariable = email, font = (36)).place(x=160,y=240)
label17 = Label(manager_add, text = "Gender:", font = (36)).place(x=10,y=270)
gender = IntVar()
radio1 = Radiobutton(manager_add, text="Male", variable = gender, value = 1, font = (36)).place(x=160,y=270)
radio2 = Radiobutton(manager_add, text="Female", variable = gender, value = 2, font = (36)).place(x=280,y=270)
label19 = Label(manager_add, text = "Type:", font = (36)).place(x=10,y=300)
role = IntVar()
radio3 = Radiobutton(manager_add, text="Manager", variable = role, value = 1, font = (36)).place(x=160,y=300)
radio4 = Radiobutton(manager_add, text="Trainer", variable = role, value = 2, font = (36)).place(x=280,y=300)
radio5 = Radiobutton(manager_add, text="Client", variable = role, value = 3, font = (36)).place(x=400,y=300)
label27 = Label(manager_add, text = "Postcode:", font = (36)).place(x=10,y=330)
postcode = StringVar()
entry7 = Entry(manager_add, textvariable = postcode, font = (36)).place(x=160,y=330)
label28 = Label(manager_add, text = "Address:", font = (36)).place(x=10,y=360)
address = StringVar()
entry8 = Entry(manager_add, textvariable = address, font = (36)).place(x=160,y=360)
submit_person_button = HoverButton(manager_add, text = "Submit",height = 2, width = 25,font = (100), command = submit_person, bg = 'black', fg = 'white').place(x=10,y=500)
clear_all_button = HoverButton(manager_add, text = "Clear All",height = 2, width = 25,font = (100), command = clear_all, bg = 'black', fg = 'white').place(x=300,y=500)
label231 = Label(manager_add, text = "Applicable to managers and trainers only:", font = (36)).place(x=780,y=130)
label232 = Label(manager_add, text = "Applicable to clients only:", font = (36)).place(x=790,y=320)
label21 = Label(manager_add, text = "Qualifications:", font = (36)).place(x=685,y=173)
text1 = tk.Text(manager_add, height=2, width=50)
text1.place(x=820,y=170)
label22 = Label(manager_add, text = "Past Experience:", font = (36)).place(x=660,y=223)
text2 = tk.Text(manager_add, height=2, width=50)
text2.place(x=820,y=220)
label233 = Label(manager_add, text = "Any underlying illnesses:", font = (36)).place(x=645,y=360)
text3 = tk.Text(manager_add, height=4, width=45)
text3.place(x=870,y=360)
label234 = Label(manager_add, text = "Any recent/affecting injuries:", font = (36)).place(x=610,y=440)
text4 = tk.Text(manager_add, height=4, width=45)
text4.place(x=870,y=440)
label24 = Label(manager_add, text = "Height/cm:", font = (36)).place(x=610,y=550)
height = StringVar()
entry5 = Entry(manager_add, textvariable = height, font = (36)).place(x=720,y=550)
label25 = Label(manager_add, text = "Weight/kg:", font = (36)).place(x=610,y=590)
weight = StringVar()
entry6 = Entry(manager_add, textvariable = weight, font = (36)).place(x=720,y=590)
smoke = IntVar()
radio6 = Radiobutton(manager_add, text="I don't smoke", variable = smoke, value = 1, font = (36)).place(x=980,y=530)
radio7 = Radiobutton(manager_add, text="I smoke", variable = smoke, value = 2, font = (36)).place(x=1150,y=530)
label26 = Label(manager_add, text = "Rate your current fitness level\n on a scale of 1 to 10:", font = (36)).place(x=980,y=570)
fitness = StringVar()
list1 = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10']
droplist = OptionMenu(manager_add,fitness,*list1)
droplist.config(width = 5)
fitness.set('1')
droplist.place(x=1090,y=625)
label27 = Label(manager_add, text = "Username:", font = (36)).place(x=10,y=390)
username_new = StringVar()
entry7 = Entry(manager_add, textvariable = username_new, font = (36)).place(x=160,y=390)
label281 = Label(manager_add, text = "Password:", font = (36)).place(x=10,y=420)
password_new = StringVar()
entry8 = Entry(manager_add, textvariable = password_new, font = (36)).place(x=160,y=420)
edit_person_button = HoverButton(manager_add, text = "Edit a person",height = 1, width = 15,font = (100), command = search_people, bg = 'black', fg = 'white').place(x=1000,y=60)

##################################################################### search ############################################################
label31 = Label(search, text = "Search for a person to edit", font = ('bold',36)).place(x=550,y=30)
treeview = ttk.Treeview(search)
treeview.config(columns = ('First name''Last name','ID'))
treeview.place(x=50,y=150)
treeview.bind('<ButtonRelease-1>',select_person)
treeview.column('#0', width = 400)
treeview.heading('#0', text = 'First name')
treeview.heading('#1', text = 'Last name')
treeview.column('#1', width = 400)
treeview.heading('#2', text = 'ID')
treeview.column('#2', width = 400)
records = treeview.get_children()
mycursor = mydb.cursor()
mycursor.execute("SELECT * FROM clients")
clients = mycursor.fetchall()
mycursor = mydb.cursor()
mycursor.execute("SELECT * FROM managers")
managers = mycursor.fetchall()
mycursor = mydb.cursor()
mycursor.execute("SELECT * FROM trainers")
trainers = mycursor.fetchall()
for i in (clients,managers,trainers):
    for row in i:
        treeview.insert('','end',text = row[1],values=(row[2],row[0]))
labe123 = Label(search, text = "search by first name:",font = (30)).place(x=10,y=10)
firstname_search = StringVar()
firstname_search_entry = Entry(search, textvariable = firstname_search, font = (30))
firstname_search_entry.place(x=210,y=13)
labe124 = Label(search, text = "search by last name:",font = (30)).place(x=10,y=50)
lastname_search = StringVar()
lastname_search_entry = Entry(search, textvariable = lastname_search, font = (30))
lastname_search_entry.place(x=210,y=53)
labe125 = Label(search, text = "search by ID:",font = (30)).place(x=10,y=90)
ID_search = StringVar()
ID_search_entry = Entry(search, textvariable = ID_search, font = (30))
ID_search_entry.place(x=210,y=93)
cancel_button1 = HoverButton(search, text = "Cancel", command = cancel_search, bg = 'black', fg = 'white', height = 1, width = 10).place(x=10,y=390)
keyboard.on_press(refine_search2, suppress=False)


################################ manager_info ###########################################
label = Label(manager_info, text = "Messages:", font = ('bold',36)).place(x=0,y=50)
manager_add_button2 = HoverButton(manager_info, text = "add People", command = manager_add_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
manager_info_button2 = HoverButton(manager_info, text = "information board", command = manager_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
manager_report_button2 = HoverButton(manager_info, text = "data reports", command = manager_report_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button2 = HoverButton(manager_info, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)


treeview11 = ttk.Treeview(manager_info,show=["headings"])
treeview11.config(columns = ('sender ID','sender name','message'),height = 15)
treeview11.place(x=10,y=130)
treeview11.bind('<ButtonRelease-1>',delete_message)
column_width = 150
treeview11.column('#1', width = column_width)
treeview11.heading('#1', text = 'sender ID')
treeview11.heading('#2', text = 'sender name')
treeview11.column('#2', width = column_width)
treeview11.heading('#3', text = 'message')
treeview11.column('#3', width = 940)
scroll_messages = ttk.Scrollbar(manager_info, orient="vertical", command=treeview11.yview)
scroll_messages.place(x=1235, y=131, height=325)
labe1 = Label(manager_info, text = "create message:",font = (30)).place(x=10,y=470)
text_message = tk.Text(manager_info, height=6, width=130)
text_message.place(x=200,y=470)
button = HoverButton(manager_info, text = "send message", command = send_message, bg = 'black', fg = 'white').place(x=10,y=520)

################################ manager_report #########################################
label6 = Label(manager_report, text = "generate reports", font = ('bold',36)).place(x=10,y=50)
manager_add_button3 = HoverButton(manager_report, text = "add People", command = manager_add_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
manager_info_button3 = HoverButton(manager_report, text = "information board", command = manager_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
manager_report_button3 = HoverButton(manager_report, text = "data reports", command = manager_report_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button3 = HoverButton(manager_report, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)
canvas = Canvas(manager_report,bg = 'light grey', width=780, height=500).place(x=420,y=100)
button = HoverButton(manager_report, text = "show graph", command = show_graph2, bg = 'black', fg = 'white', height = 2, width = 20, font = 30).place(x=10,y=470)
button = HoverButton(manager_report, text = "save graph", command = gen_report2, bg = 'black', fg = 'white', height = 2, width = 20, font = 30).place(x=10,y=540)


treeview14 = ttk.Treeview(manager_report)
treeview14.config(columns = ('report'), height = 15)
treeview14.place(x=10,y=130)
treeview14.bind('<ButtonRelease-1>',select_exercise)
vsb = ttk.Scrollbar(manager_report, orient="vertical", command=treeview14.yview)
vsb.place(x=275, y=131, height=325)
treeview14.configure(yscrollcommand=vsb.set)
treeview14.column('#0', width = 280)
treeview14.heading('#0', text = 'report')
treeview14.column('#1', width = 0)
treeview14.heading('#1', text = '')


################################ trainer_make_classes ###################################
label7 = Label(trainer_make_classes, text = "create a class:", font = ('bold',36)).place(x=10,y=50)
trainer_make_button1 = HoverButton(trainer_make_classes, text = "make classes", command = trainer_make_classes_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
trainer_info_button1 = HoverButton(trainer_make_classes, text = "information board", command = trainer_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
trainer_see_bookings_button1 = HoverButton(trainer_make_classes, text = "see bookings", command = trainer_see_bookings_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button4 = HoverButton(trainer_make_classes, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)
labe1 = Label(trainer_make_classes, text = "select start time:",font = (30)).place(x=10,y=180)
labe1 = Label(trainer_make_classes, text = ":",font = (30)).place(x=100,y=220)
time1 = StringVar()
list_time1 = []
for i in range(1,25):
    list_time1.append(str(i))
droplist = OptionMenu(trainer_make_classes,time1,*list_time1)
droplist.config(width = 5)
time1.set('1')
droplist.place(x=10,y=220)
time2 = StringVar()
list_time2 = []
for i in range(0,61):
    list_time2.append(str(i))
droplist = OptionMenu(trainer_make_classes,time2,*list_time2)
droplist.config(width = 5)
time2.set('1')
droplist.place(x=125,y=220)
labe1 = Label(trainer_make_classes, text = "class duration:",font = (30)).place(x=10,y=260)
class_duration = StringVar()
entry = Entry(trainer_make_classes, textvariable = class_duration, font = (36)).place(x=10,y=300)
labe1 = Label(trainer_make_classes, text = "mins",font = (30)).place(x=250,y=298)
monday = IntVar()
tuesday = IntVar()
wednesday = IntVar()
thursday = IntVar()
friday = IntVar()
saturday = IntVar()
sunday = IntVar()
Checkbutton(trainer_make_classes, text="monday", variable=monday, font = 80).place(x=10,y=350)
Checkbutton(trainer_make_classes, text="tuesday", variable=tuesday, font = 80).place(x=120,y=350)
Checkbutton(trainer_make_classes, text="wednesday", variable=wednesday, font = 80).place(x=230,y=350)
Checkbutton(trainer_make_classes, text="thursday", variable=thursday, font = 80).place(x=370,y=350)
Checkbutton(trainer_make_classes, text="friday", variable=friday, font = 80).place(x=480,y=350)
Checkbutton(trainer_make_classes, text="saturday", variable=saturday, font = 80).place(x=560,y=350)
Checkbutton(trainer_make_classes, text="sunday", variable=sunday, font = 80).place(x=670,y=350)
labe1 = Label(trainer_make_classes, text = "class overview:",font = (30)).place(x=10,y=420)
text984 = tk.Text(trainer_make_classes, height=8, width=100)
text984.place(x=10,y=450)
button = HoverButton(trainer_make_classes, text = "create class", command = create_class, bg = 'black', fg = 'white', height = 3, width = 35, font = 30).place(x=850,y=470)

treeview8 = ttk.Treeview(trainer_make_classes)
treeview8.config(columns = ('class ID','days''start time','end time','overview'))
treeview8.place(x=340,y=110)
treeview8.bind('<ButtonRelease-1>',select_class)
treeview8.heading('#0', text = 'class ID')
treeview8.column('#0', width = 140)
treeview8.heading('#1', text = 'days')
treeview8.column('#1', width = 160)
treeview8.heading('#2', text = 'start time')
treeview8.column('#2', width = 150)
treeview8.heading('#3', text = 'end time')
treeview8.column('#3', width = 150)
treeview8.heading('#4', text = 'overview')
treeview8.column('#4', width = 320)
scroll_trainer_make = ttk.Scrollbar(trainer_make_classes, orient="vertical", command=treeview8.yview)
scroll_trainer_make.place(x=1247, y=111, height=225)
records = treeview8.get_children()



################################ trainer_info ###########################################
label = Label(trainer_info, text = "Messages:", font = ('bold',36)).place(x=0,y=50)
trainer_make_button2 = HoverButton(trainer_info, text = "make classes", command = trainer_make_classes_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
trainer_info_button2 = HoverButton(trainer_info, text = "information board", command = trainer_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
trainer_see_bookings_button2 = HoverButton(trainer_info, text = "see bookings", command = trainer_see_bookings_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button5 = HoverButton(trainer_info, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)

treeview12 = ttk.Treeview(trainer_info,show=["headings"])
treeview12.config(columns = ('sender ID','sender name','message'),height = 15)
treeview12.place(x=10,y=130)
treeview12.bind('<ButtonRelease-1>',delete_message2)
column_width = 150
treeview12.column('#1', width = column_width)
treeview12.heading('#1', text = 'sender ID')
treeview12.heading('#2', text = 'sender name')
treeview12.column('#2', width = column_width)
treeview12.heading('#3', text = 'message')
treeview12.column('#3', width = 940)
scroll_messages = ttk.Scrollbar(trainer_info, orient="vertical", command=treeview12.yview)
scroll_messages.place(x=1235, y=131, height=325)
labe1 = Label(trainer_info, text = "create message:",font = (30)).place(x=10,y=470)
text_message2 = tk.Text(trainer_info, height=6, width=130)
text_message2.place(x=200,y=470)
button = HoverButton(trainer_info, text = "send message", command = send_message2, bg = 'black', fg = 'white', height = 1, width = 15, font = 30).place(x=10,y=520)

################################ trainer_see_bookings ###################################
label9 = Label(trainer_see_bookings, text = "upcoming classes:", font = ('bold',36)).place(x=10,y=50)
trainer_make_button3 = HoverButton(trainer_see_bookings, text = "make classes", command = trainer_make_classes_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
trainer_info_button3 = HoverButton(trainer_see_bookings, text = "information board", command = trainer_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
trainer_see_bookings_button3 = HoverButton(trainer_see_bookings, text = "see bookings", command = trainer_see_bookings_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button6 = HoverButton(trainer_see_bookings, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)


treeview9 = ttk.Treeview(trainer_see_bookings,show=["headings"])
treeview9.config(columns = ('date','class ID','start time','end time','overview'),height = 8)
treeview9.place(x=10,y=130)
treeview9.bind('<ButtonRelease-1>',check_bookings)
column_width = 150
treeview9.column('#1', width = column_width)
treeview9.heading('#1', text = 'date')
treeview9.heading('#2', text = 'class ID')
treeview9.column('#2', width = column_width)
treeview9.heading('#3', text = 'start time')
treeview9.column('#3', width = column_width)
treeview9.heading('#4', text = 'end time')
treeview9.column('#4', width = column_width)
treeview9.heading('#5', text = 'overview')
treeview9.column('#5', width = 300)
scroll_bookings = ttk.Scrollbar(trainer_see_bookings, orient="vertical", command=treeview9.yview)
scroll_bookings.place(x=896, y=131, height=185)

treeview10 = ttk.Treeview(trainer_see_bookings,show=["headings"])
treeview10.config(columns = ('client ID','class ID','date','name','phone'),height = 8)
treeview10.place(x=10,y=330)
treeview10.bind('<ButtonRelease-1>',mark_as_present)
column_width = 150
treeview10.column('#1', width = column_width)
treeview10.heading('#1', text = 'client ID')
treeview10.heading('#2', text = 'class ID')
treeview10.column('#2', width = column_width)
treeview10.heading('#3', text = 'date')
treeview10.column('#3', width = column_width)
treeview10.heading('#4', text = 'name')
treeview10.column('#4', width = column_width)
treeview10.heading('#5', text = 'phone')
treeview10.column('#5', width = 300)
scroll_bookings = ttk.Scrollbar(trainer_see_bookings, orient="vertical", command=treeview9.yview)
scroll_bookings.place(x=896, y=331, height=185)

################################ client_book ############################################
label = Label(client_book, text = "Search for a class to book:", font = ('bold',30)).place(x=10,y=45)
label = Label(client_book, text = "my bookings:", font = ('bold',30)).place(x=10,y=345)
cal2 = Calendar(client_book,bg="darkblue",fg="white",date_pattern = 'dd/mm/yy',mindate = date.today()+ timedelta(days=1))
cal2.place(x=20,y=118)
cal2.bind('<<CalendarSelected>>',update_classes)
treeview3 = ttk.Treeview(client_book,show=["headings"])
treeview3.config(columns = ('trainer','start time','end time','date','overview'))
treeview3.place(x=300,y=100)
treeview3.bind('<ButtonRelease-1>',book_class)
vsb2 = ttk.Scrollbar(client_book, orient="vertical", command=treeview3.yview)
vsb2.place(x=1203, y=101, height=225)
column_width = 150
treeview3.column('#1', width = column_width)
treeview3.heading('#1', text = 'trainer')
treeview3.heading('#2', text = 'start time')
treeview3.column('#2', width = column_width)
treeview3.heading('#3', text = 'end time')
treeview3.column('#3', width = column_width)
treeview3.heading('#4', text = 'date')
treeview3.column('#4', width = column_width)
treeview3.heading('#5', text = 'overview')
treeview3.column('#5', width = 300)
records = treeview3.get_children()
treeview4 = ttk.Treeview(client_book,show=["headings"])
treeview4.config(columns = ('trainer','start time','end time','date','overview'))
treeview4.place(x=300,y=400)
treeview4.bind('<ButtonRelease-1>',delete_selected_booking)
vsb3 = ttk.Scrollbar(client_book, orient="vertical", command=treeview4.yview)
vsb3.place(x=1203, y=401, height=225)
column_width = 150
treeview4.heading('#1', text = 'trainer')
treeview4.column('#1', width = column_width)
treeview4.heading('#2', text = 'start time')
treeview4.column('#2', width = column_width)
treeview4.heading('#3', text = 'end time')
treeview4.column('#3', width = column_width)
treeview4.heading('#4', text = 'date')
treeview4.column('#4', width = column_width)
treeview4.heading('#5', text = 'overview')
treeview4.column('#5', width = 300)
records = treeview3.get_children()
client_book_button1 = HoverButton(client_book, text = "book classes", command = client_book_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
client_info_button1 = HoverButton(client_book, text = "information board", command = client_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
client_track_button1 = HoverButton(client_book, text = "track fitness", command = client_track_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button7 = HoverButton(client_book, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)
label = Label(client_book, text = " Alert me             minutes \n before my class using:", font = (36)).place(x=10,y=420)
alert_time = IntVar()
list2 = [5,10,15,20,25,30,35,40,45,50,55,60]
droplist2 = OptionMenu(client_book,alert_time,*list2)
droplist2.config(width = 3)
droplist2.place(x=100,y=420)
message_type = IntVar()
radio = Radiobutton(client_book, text="SMS", variable = message_type, value = 1, font = (36)).place(x=10,y=480)
radio = Radiobutton(client_book, text="Email", variable = message_type, value = 2, font = (36)).place(x=120,y=480)
radio = Radiobutton(client_book, text="Both", variable = message_type, value = 3, font = (36)).place(x=10,y=530)
radio = Radiobutton(client_book, text="Neither", variable = message_type, value = 4, font = (36)).place(x=120,y=530)


################################ client_info ############################################
label = Label(client_info, text = "Messages:", font = ('bold',36)).place(x=0,y=50)
client_book_button2 = HoverButton(client_info, text = "book classes", command = client_book_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
client_info_button2 = HoverButton(client_info, text = "information board", command = client_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
client_track_button2 = HoverButton(client_info, text = "track fitness", command = client_track_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button8 = HoverButton(client_info, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)
client_map_button = HoverButton(client_info, text = "show map", command = show_map, bg = 'black', fg = 'white', height = 2, width = 40).place(x=850,y=550)

treeview13 = ttk.Treeview(client_info,show=["headings"])
treeview13.config(columns = ('sender ID','sender name','message'),height = 18)
treeview13.place(x=10,y=130)
treeview13.bind('<ButtonRelease-1>',view_message)
column_width = 150
treeview13.column('#1', width = column_width)
treeview13.heading('#1', text = 'sender ID')
treeview13.heading('#2', text = 'sender name')
treeview13.column('#2', width = column_width)
treeview13.heading('#3', text = 'message')
treeview13.column('#3', width = 940)
scroll_messages = ttk.Scrollbar(client_info, orient="vertical", command=treeview13.yview)
scroll_messages.place(x=1235, y=131, height=385)

message_alert_type = IntVar()
radio = Radiobutton(client_info, text="SMS", variable = message_alert_type, value = 1, font = (36)).place(x=290,y=550)
radio = Radiobutton(client_info, text="Email", variable = message_alert_type, value = 2, font = (36)).place(x=400,y=550)
radio = Radiobutton(client_info, text="Both", variable = message_alert_type, value = 3, font = (36)).place(x=510,y=550)
radio = Radiobutton(client_info, text="Neither", variable = message_alert_type, value = 4, font = (36)).place(x=620,y=550)
label = Label(client_info, text = " Alert me to messages using:", font = (36)).place(x=10,y=550)

################################ client_track ###########################################
canvas8 = Canvas(client_track,bg = 'light grey', width=900, height=900).place(x=480,y=0)
client_book_button3 = HoverButton(client_track, text = "book classes", command = client_book_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=0,y=0)
client_info_button3 = HoverButton(client_track, text = "information board", command = client_info_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=430,y=0)
client_track_button3 = HoverButton(client_track, text = "track fitness", command = client_track_goto, bg = 'black', fg = 'white', height = 2, width = 60).place(x=860,y=0)
logout_button9 = HoverButton(client_track, text = "Logout", command = logout, bg = 'black', fg = 'white').place(x=10,y=620)
new_exercise = StringVar()
new_exercise_entry = Entry(client_track, textvariable = new_exercise, font = (30)).place(x=10,y=53)
submit_new_exercise_button = HoverButton(client_track, text = "create a new exercise", command = submit_new_exercise, bg = 'black', fg = 'white').place(x=250,y=53)
treeview2 = ttk.Treeview(client_track)
treeview2.config(columns = ('exercise'))
treeview2.place(x=10,y=130)
treeview2.bind('<ButtonRelease-1>',select_exercise)
vsb = ttk.Scrollbar(client_track, orient="vertical", command=treeview2.yview)
vsb.place(x=196, y=131, height=225)
treeview2.configure(yscrollcommand=vsb.set)
treeview2.column('#0', width = 200)
treeview2.heading('#0', text = 'exercise')
treeview2.column('#1', width = 0)
treeview2.heading('#1', text = '')
records2 = treeview2.get_children()
delete_selected_exercise_button = HoverButton(client_track, text = "delete the selected exercise", command = delete_selected_exercise, bg = 'black', fg = 'white').place(x=250,y=150)
label15 = Label(client_track, text = "weight/kg:", font = (36)).place(x=10,y=400)
weight2 = StringVar()
entry1 = Entry(client_track, textvariable = weight2, font = (36)).place(x=110,y=403)
label15 = Label(client_track, text = "reps:", font = (36)).place(x=10,y=430)
reps = StringVar()
entry1 = Entry(client_track, textvariable = reps, font = (36)).place(x=110,y=433)
label15 = Label(client_track, text = "sets:", font = (36)).place(x=10,y=460)
sets = StringVar()
entry1 = Entry(client_track, textvariable = sets, font = (36)).place(x=110,y=463)
label15 = Label(client_track, text = "time/s:", font = (36)).place(x=10,y=490)
sesh_time = StringVar()
entry1 = Entry(client_track, textvariable = sesh_time, font = (36)).place(x=110,y=493)
add_data_button = HoverButton(client_track, text = "record your performance", command = add_data, bg = 'black', fg = 'white').place(x=250,y=200)
show_graph_button = HoverButton(client_track, text = "show graph", command = show_graph_exercise, bg = 'black', fg = 'white').place(x=250,y=250)
report_button = HoverButton(client_track, text = "generate report", command = generate_report, bg = 'black', fg = 'white').place(x=250,y=300)

#########################################################################################
DB_fill()
role.set(3)
smoke.set(1)
gender.set(1)
send_alert()
raise_frame(login)
root.mainloop()


#root.after(2000,printing)
#INSERT INTO logins(userid,username,password_) VALUES ('m1','test1','123');
#CREATE TABLE trainers( ID VARCHAR(255), firstname VARCHAR(255), lastname VARCHAR(255), DOB VARCHAR(255), phone VARCHAR(255), email VARCHAR(255), gender VARCHAR(255), postcode VARCHAR(255), address VARCHAR(255), qualifications VARCHAR(255), experience VARCHAR(255));
